# -*- coding:utf-8 -*-
# import os
# import docx
#
# project_path = os.path.dirname(os.path.abspath(__file__))
# file_path = os.path.join(project_path,'python_exercise_100','沈阳桃李面包股份有限公司.docx')
# file = docx.Document(file_path)

import os.path
from docx import Document
path="C:\\Users\\xiong\\Desktop\\qiuxiaoshuang1\\python_exercise_100\\沈阳桃李面包股份有限公司.docx"
project_path = os.path.dirname(os.path.abspath(__file__))
print(project_path)
file_path = os.path.join(project_path,'沈阳桃李面包股份有限公司.docx')
document=Document(file_path)
for paragraph in document.paragraphs:
    print(paragraph.text)