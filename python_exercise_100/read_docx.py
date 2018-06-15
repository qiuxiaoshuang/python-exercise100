# -*- coding:utf-8 -*-
import os
import docx

project_path = os.path.dirname(os.path.abspath(__file__))
file_path = os.path.join(project_path,'沈阳桃李面包股份有限公司.docx')
print(file_path)
file = docx.Document(file_path)
para_nums = len(file.paragraphs)

print("段落数： " + str(para_nums))



# for para in file.paragraphs:
#     print(para.text)

def is_words_in_paragraph(words,paragraph):
    if words in paragraph:
        return True
    else:
        return False

# words='首次公开发行股票招股说明书'
#
# for i in range(para_nums):
#     if len(file.paragraphs[i].text) > 0 :
#         para = file.paragraphs[i].text
#         if is_words_in_paragraph(words,para):
#             print("***********")
#         print("第 " + str(i)+" 段的内容是： ")
#         print(para)



for table in file.tables:
    print("********** table **********")
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)



